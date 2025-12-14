import os
import re
import subprocess
import json
import fitz  
import pandas as pd
import openpyxl
import jieba
import logging
import time
import argparse  
from datetime import datetime
from collections import Counter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import jieba.posseg as pseg
from tqdm import tqdm


# 导入配置 (注意：你需要确保 src 在 pythonpath 中，后面会讲怎么运行)
from src.settings import RAW_DATA_PATH, PROCESSED_DATA_PATH, FAIL_DATA_PATH,LOG_DIR, STOPWORDS_PATH
INPUT_DATA = RAW_DATA_PATH / 'data.json'  
OUTPUT_JSON = PROCESSED_DATA_PATH / 'extract_data.json'
FAIL_JSON = FAIL_DATA_PATH / 'fail.json'


# =========================================================================
# 组件 0: 日志系统配置
# =========================================================================

def setup_logger(log_file_path):
    """
    配置双向日志：
    1. 终端显示 (StreamHandler) - 实际由 tqdm.write 接管
    2. 文件记录 (FileHandler)
    """
    # 创建 Logger
    logger = logging.getLogger("doc_pipeline")
    logger.setLevel(logging.INFO)
    logger.handlers = []  # 清除已有 handler 避免重复

    # 格式器
    formatter = logging.Formatter('%(asctime)s %(message)s', datefmt='%Y-%m-%d %H:%M:%S')

    # 1. 文件处理器
    file_handler = logging.FileHandler(log_file_path, encoding='utf-8')
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)

    return logger

# 全局 logger
logger = None

def log_msg(level, tag, msg):
    """
    统一日志打印函数，兼容 tqdm 和 logging 文件保存
    level: 'info', 'error'
    tag: '[INFO]', '[Success]', '[FAIL]'
    """
    full_msg = f"{tag} {msg}"
    
    # 写入日志文件
    if logger:
        if level == 'error':
            logger.error(full_msg)
        else:
            logger.info(full_msg)
            
    # 终端显示 (使用 tqdm.write 防止进度条错位)
    tqdm.write(full_msg)


# =========================================================================
# 组件 1: 文本分词器 (NLP 处理层)
# =========================================================================

class TextTokenizer:
    def __init__(self, stop_words_path=None):
        self.stop_words = self._load_default_stopwords()
        if stop_words_path and os.path.exists(stop_words_path):
            try:
                with open(stop_words_path, 'r', encoding='utf-8') as f:
                    external_words = set(line.strip() for line in f)
                    self.stop_words.update(external_words)
            except Exception as e:
                log_msg('error', '[FAIL]', f"加载停用词文件失败: {e}")

    def _load_default_stopwords(self):
        words = {
            "的", "了", "和", "是", "就", "都", "而", "及", "与", "着", "或", 
            "一个", "没有", "我们", "你们", "他们", "它", "它们", 
            "在", "从", "对", "对于", "把", "被", "让", "向", "往", 
            "虽然", "但是", "因为", "所以", "如果", "那么", "以及", 
            "什么", "怎么", "哪里", "哪个", "这里", "那里",
            "建议", "意见", "办法", "情况", 
            "\n", "\t", " ", "\u3000", "\xa0"
        }
        return words

    def tokenize(self, text):
        if not text:
            return []
        text = re.sub(r'\s+', ' ', text)
        words = pseg.cut(text)
        valid_tokens = []
        for word, flag in words:
            word = word.strip()
            if len(word) < 2: continue
            if word in self.stop_words: continue
            if flag.startswith('m'): continue
            if flag.startswith('q'): continue
            if flag.startswith('x'): continue
            if flag.startswith('w'): continue
            if flag.startswith('p'): continue
            if flag.startswith('c'): continue
            if flag.startswith('u'): continue
            if flag.startswith('r'): continue
            if flag.startswith('t'): continue
            valid_tokens.append(word)
        return valid_tokens

# =========================================================================
# 组件 2: 文档提取器 (IO 与 解析层)
# =========================================================================

class FileContentExtractor:
    def _clean_filename_as_title(self, filename):
        name = os.path.splitext(filename)[0]
        name = re.sub(r'[-_.]', ' ', name)
        return name.strip()

    def _is_meaningful_filename(self, filename):
        name_no_ext = os.path.splitext(filename)[0]
        if len(name_no_ext) < 2: return False
        clean_text = re.sub(r'[-_.]', ' ', name_no_ext).strip()
        if re.match(r'^\d+$', clean_text.replace(' ', '')): return False
        if re.match(r'^[\d\s/-]+$', clean_text): return False
        if re.match(r'^[\W_]+$', clean_text.replace(' ', '')): return False

        words = jieba.lcut(clean_text)
        total_len = len(clean_text.replace(' ', ''))
        if total_len == 0: return False
        semantic_len = 0
        for w in words:
            w = w.strip()
            if not w: continue
            if any('\u4e00' <= char <= '\u9fff' for char in w):
                semantic_len += len(w)
            elif w.isalpha() and len(w) > 2:
                if w.lower() in {'scan', 'img', 'doc', 'file', 'temp', 'untitled', 'screenshot'}:
                    continue
                semantic_len += len(w)
        score = semantic_len / total_len
        return score >= 0.4

    def extract(self, filepath):
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"文件路径不存在: {filepath}")

        ext = os.path.splitext(filepath)[1].lower()
        segments, title = [], ""
        file_basename = os.path.basename(filepath)
        filename_is_good = self._is_meaningful_filename(file_basename)
        clean_filename = self._clean_filename_as_title(file_basename)

        try:
            if ext == '.pdf':
                segments, title = self._process_pdf(filepath)
            elif ext == '.docx':
                segments, title = self._process_docx(filepath)
            elif ext == '.doc':
                segments, title = self._process_doc_legacy(filepath)
            elif ext in ['.xlsx', '.xls']:
                segments, title = self._process_excel(filepath)
            elif ext == '.txt':
                segments, title = self._process_txt(filepath)
            else:
                raise ValueError(f"不支持的文件格式: {ext}")
            
            raw_text = self._smart_merge_segments(segments)
            final_content = self._remove_noise_chars(raw_text)
            extra_title = self._remove_noise_chars(title)
            
            final_title = ""
            if filename_is_good:
                final_title = clean_filename
            else:
                if extra_title and len(extra_title.strip()) > 1:
                    final_title = extra_title
            
            return final_content, final_title

        except Exception as e:
            raise RuntimeError(f"解析过程异常 ({ext}): {str(e)}")

    # --- 核心工具方法 ---

    @staticmethod
    def _is_cjk_char(char):
        if not char: return False
        code = ord(char)
        return (0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF or 0x20000 <= code <= 0x2A6DF)

    @staticmethod
    def _remove_noise_chars(text):
        if not text: return ""
        text = re.sub(r'_{2,}', ' ', text)
        text = re.sub(r'\.{3,}', ' ', text)
        text = text.replace('（', '(').replace('）', ')')
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    @staticmethod
    def _smart_merge_segments(text_segments):
        if not text_segments: return ""
        buffer = []
        last_char = None
        for segment in text_segments:
            clean_seg = re.sub(r'\s+', ' ', segment.strip())
            if not clean_seg: continue
            if not buffer:
                buffer.append(clean_seg)
                last_char = clean_seg[-1]
                continue
            first_char = clean_seg[0]
            if last_char == '-' and 'a' <= first_char.lower() <= 'z':
                buffer[-1] = buffer[-1].rstrip('-')
                buffer.append(clean_seg)
            elif (FileContentExtractor._is_cjk_char(last_char) and 
                  FileContentExtractor._is_cjk_char(first_char)):
                buffer.append(clean_seg)
            else:
                buffer.append(" " + clean_seg)
            last_char = clean_seg[-1]
        return "".join(buffer)

    # --- 各格式处理逻辑 ---

    def _process_txt(self, filepath):
        segments, title = [], ""
        content = ""
        for enc in ['utf-8', 'gb18030', 'gbk', 'utf-16', 'latin-1']:
            try:
                with open(filepath, 'r', encoding=enc) as f: content = f.read()
                break
            except: continue
        
        if content: 
            segments = content.splitlines()
            for line in segments:
                clean_line = line.strip()
                if clean_line:
                    title = clean_line
                    break
        else:
            raise ValueError("TXT读取内容为空或编码不支持")

        return segments, title

    def _process_pdf(self, filepath):
        segments = []
        try:
            doc = fitz.open(filepath)
        except Exception as e:
            raise ValueError(f"PDF文件损坏或无法打开: {e}")

        title = self._get_pdf_title_optimized(doc)
        if not title and "title" in doc.metadata: title = doc.metadata["title"]
        
        for page in doc:
            h = page.rect.height
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))
            for b in blocks:
                if b[6] != 0: continue
                if b[3] < h * 0.08 or b[1] > h * 0.92: continue
                txt = b[4].strip()
                if txt: segments.append(txt.replace('\n', ' '))
        doc.close()
        return segments, title

    def _get_pdf_title_optimized(self, doc):
        try:
            page = doc[0]
            lines_data = []
            for block in page.get_text("dict")["blocks"]:
                if "lines" not in block: continue
                for line in block["lines"]:
                    line_text = "".join([s["text"] for s in line["spans"]]).strip()
                    max_size = max([s["size"] for s in line["spans"]]) if line["spans"] else 0
                    if line_text and max_size > 0:
                        lines_data.append({"text": line_text, "size": max_size, "y": line["bbox"][1]})
            if not lines_data: return ""
            max_size = max(x["size"] for x in lines_data)
            limit_y, limit_size = page.rect.height * 0.35, max_size * 0.8
            parts = [x["text"] for x in sorted(lines_data, key=lambda i: i["y"]) 
                     if x["y"] < limit_y and x["size"] >= limit_size]
            return " ".join(parts) if parts else (max(lines_data, key=lambda x:x["size"])["text"] if lines_data else "")
        except: return ""

    def _process_docx(self, filepath):
        segments = []
        try:
            doc = Document(filepath)
        except Exception as e:
            raise ValueError(f"DOCX文件损坏: {e}")

        title = self._get_docx_title_optimized(doc)
        if not title and doc.core_properties.title: title = doc.core_properties.title
        for p in doc.paragraphs:
            if p.text.strip(): segments.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                cells = [c.text.strip() for c in r.cells if c.text.strip()]
                if cells: segments.append(" ".join(cells))
        return segments, title

    def _get_docx_title_optimized(self, doc):
        parts, scan_limit, prev_is_title = [], 8, False
        for i, para in enumerate(doc.paragraphs):
            if i >= scan_limit: break
            text = para.text.strip()
            if not text: continue
            score = 0
            style = para.style.name.lower()
            if 'title' in style: score += 100
            elif 'heading 1' in style: score += 50
            max_pt = 0
            for run in para.runs:
                if run.font.size and run.font.size.pt > max_pt: max_pt = run.font.size.pt
            if max_pt == 0: max_pt = 16 if score >= 50 else 10.5
            if max_pt >= 15: score += 40
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER: score += 20
            if score > 40:
                parts.append(text)
                prev_is_title = True
            elif prev_is_title: break
        return " ".join(parts) if parts else ""

    def _process_excel(self, filepath):
        segments, title = [], ""
        xls = pd.ExcelFile(filepath) 
        for i, sheet in enumerate(xls.sheet_names):
            df = pd.read_excel(xls, sheet_name=sheet, header=None).fillna("").astype(str)
            if i == 0 and not title and not df.empty:
                row0 = [x.strip() for x in df.iloc[0] if x.strip() and x.lower()!='nan']
                if len(row0) == 1: title = row0[0]
            if len(xls.sheet_names)>1: segments.append(f"Sheet: {sheet}")
            for _, row in df.iterrows():
                vals = [x.strip() for x in row if x.strip() and x.lower()!='nan']
                if vals: segments.append(" ".join(vals))
        return segments, title

    def _process_doc_legacy(self, filepath):
        res = subprocess.run(['antiword', '-w', '0', filepath], capture_output=True, text=True)
        if res.returncode == 0 and res.stdout:
            lines = res.stdout.splitlines()
            title = ""
            for line in lines:
                if line.strip():
                    title = line.strip()
                    break
            return lines, title
        else:
            raise ValueError(f"Antiword 执行失败 (return code {res.returncode})")

# =========================================================================
# 组件 3: 处理流水线 (Pipeline)
# =========================================================================

class DocumentPipeline:
    def __init__(self, stop_words_path=None):
        self.extractor = FileContentExtractor() 
        self.tokenizer = TextTokenizer(stop_words_path=stop_words_path)

    def run(self, filepath):
        """
        Pipeline 主流程
        """
        content, title = self.extractor.extract(filepath)
        
        if not content.strip():
            return None 

        seg_title = self.tokenizer.tokenize(title)
        seg_content = self.tokenizer.tokenize(content)

        return {
            "title": title,
            "content": content,
            "seg_title": seg_title,
            "seg_content": seg_content,
        }

# =========================================================================
# 测试代码
# =========================================================================

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Document Extraction Pipeline")
    parser.add_argument(
        '--mode', 
        type=int, 
        default=10, 
        help="Number of files to process. Set 0 to process all files. (Default: 0)"
    )
    args = parser.parse_args()

    # 配置

    # 生成带时间戳的 Log 文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    LOG_FILE = os.path.join(LOG_DIR, f'extract_data_{timestamp}.log')

    # 初始化日志
    logger = setup_logger(LOG_FILE)
    log_msg('info', '[INFO]', f"Pipeline initialized. Logs saving to: {os.path.abspath(LOG_FILE)}")

    # [修改] 模式选择：使用命令行参数 args.mode
    mode = args.mode
    log_msg('info', '[INFO]', f"Mode set to: {mode} (Processing {'ALL files' if mode == 0 else f'first {mode} files'})")

    if isinstance(mode, int):
        VALID_EXTS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt'}

        log_msg('info', '[INFO]', "Loading and deduplicating data...")
        try:
            with open(INPUT_DATA, 'r', encoding='utf-8') as f:
                if mode:
                    data_list = json.load(f)[:mode]
                else:
                    data_list = json.load(f)
        except Exception as e:
            log_msg('error', '[FAIL]', f"Failed to load input data: {e}")
            exit(1)

        # 去重逻辑
        seen_keys = set()
        unique_data_list = []
        total_raw = len(data_list)
        
        for item in data_list:
            key = item.get('url') or item.get('path')
            if key and key not in seen_keys:
                seen_keys.add(key)
                unique_data_list.append(item)
        
        data_list = unique_data_list
        log_msg('info', '[INFO]', f"Deduplication complete. Raw: {total_raw} -> Unique: {len(data_list)}")

        log_msg('info', '[INFO]', "Initializing DocumentPipeline...")
        pipeline = DocumentPipeline(STOPWORDS_PATH)

        log_msg('info', '[INFO]', "Starting file processing...")
        final_results = []
        failed_results = [] 
        count = 0

        for item in tqdm(data_list, desc="Processing", unit="file"):
            relative_path = item.get('path').lstrip('/') 
            filepath = RAW_DATA_PATH / relative_path
            url = item.get('url')
            filename = os.path.basename(filepath)
            
            # 单个文件的处理闭环
            try:
                ext = os.path.splitext(filepath)[1].lower()
                if ext not in VALID_EXTS:
                    raise ValueError(f"Skipped extension: {ext}")

                result = pipeline.run(filepath)
                
                if result:
                    result['url'] = url
                    final_results.append(result)
                    
                    t_sample = ''.join(result['title'][:10]) if result.get('title') else "No Title"
                    c_sample = ' '.join(result['seg_content'][:20]) if result.get('seg_content') else ""
                    log_msg('info', '[Success]', f"{filename} | Title: {t_sample}... | Seg: {c_sample}...")
                else:
                    reason = "Extraction returned empty content"
                    item['error_reason'] = reason
                    item['failure_type'] = "EmptyContent"
                    failed_results.append(item)
                    log_msg('info', '[FAIL]', f"Skipped (Empty): {filename}")
                    
            except Exception as e:
                error_msg = str(e)
                item['error_reason'] = error_msg
                
                # 简单分类错误类型
                if "FileNotFound" in error_msg or "不存在" in error_msg:
                    item['failure_type'] = "FileNotFound"
                elif "ValueError" in error_msg:
                    item['failure_type'] = "FormatError"
                else:
                    item['failure_type'] = "RuntimeError"

                failed_results.append(item)
                log_msg('error', '[FAIL]', f"Error processing {filename}: {error_msg}")
            
            count += 1

        # 写入成功结果
        try:
            with open(OUTPUT_JSON, 'w', encoding='utf-8') as f:
                json.dump(final_results, f, ensure_ascii=False, indent=4)
            log_msg('info', '[INFO]', "--------------------------------------------------")
            log_msg('info', '[Success]', "Data extraction complete.")
            log_msg('info', '[INFO]', f"Total processed: {len(data_list)}")
            log_msg('info', '[INFO]', f"Successful: {len(final_results)} (Saved to: {os.path.abspath(OUTPUT_JSON)})")
        except Exception as e:
            log_msg('error', '[FAIL]', f"Failed to write output JSON: {e}")

        # 写入失败结果 (带原因和统计)
        if failed_results:
            try:
                # 1. 写入 fail.json
                with open(FAIL_JSON, 'w', encoding='utf-8') as f:
                    json.dump(failed_results, f, ensure_ascii=False, indent=4)
                
                # 2. 统计失败类型并打印到 Log
                fail_types = [item.get('failure_type', 'Unknown') for item in failed_results]
                type_counts = Counter(fail_types)
                
                log_msg('info', '[INFO]', f"Failed: {len(failed_results)} (Saved to: {os.path.abspath(FAIL_JSON)})")
                log_msg('info', '[INFO]', "Failure Type Breakdown:")
                for f_type, count in type_counts.items():
                    log_msg('info', '[INFO]', f"   - {f_type}: {count}")
                    
            except Exception as e:
                log_msg('error', '[FAIL]', f"Failed to write fail JSON: {e}")
        else:
            log_msg('info', '[INFO]', "Failed: 0 (No failures to record)")
            
        log_msg('info', '[INFO]', "--------------------------------------------------")